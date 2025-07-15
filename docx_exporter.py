# -*- coding: utf-8 -*-
# docx_exporter.py

from docx import Document
from docx.shared import Inches
from utils import format_number, clean_text
import io

def generate_word_report(
    data,
    ticker,
    closing_price,
    peg,
    benchmark,
    targets,
    valuation_label,
    ai_summary,
    chart_image=None,
    selected_range=None  # <-- NEW: Optional chart time range label
):
    doc = Document()

    # Title
    doc.add_heading(f"{data.get('shortName', '')} ({ticker}) Financial Report", level=0)

    # Company Overview
    doc.add_heading("Company Overview", level=1)
    doc.add_paragraph(f"Industry: {data.get('industry', 'N/A')}")
    doc.add_paragraph(f"Sector: {data.get('sector', 'N/A')}")
    doc.add_paragraph(f"Market Cap: {format_number(data.get('marketCap', 0), style='usd')}")
    doc.add_paragraph(f"Current Price: {format_number(closing_price, style='usd')}")

    # Key Financial Metrics
    doc.add_heading("Key Financial Metrics", level=1)

    metrics = {
        "Trailing PE": ("trailingPE", "ratio"),
        "Forward PE": ("forwardPE", "ratio"),
        "Return on Equity (ROE)": ("returnOnEquity", "percent"),
        "Gross Margin": ("grossMargins", "percent"),
        "Operating Margin": ("operatingMargin", "percent"),
        "Debt to Equity": ("debtToEquity", "ratio"),
        "PEG Ratio (from Yahoo)": (peg, "ratio")  # PEG is passed as a value
    }

    for label, (key_or_value, style) in metrics.items():
        paragraph = doc.add_paragraph()
        run_label = paragraph.add_run(f"{label}: ")
        run_label.bold = True

        if label == "PEG Ratio (from Yahoo)":
            formatted_value = format_number(key_or_value, style)
        else:
            raw_value = data.get(key_or_value)
            formatted_value = format_number(raw_value, style)

        paragraph.add_run(formatted_value)

    # Sector Benchmark Comparison
    if benchmark:
        doc.add_heading("Sector Benchmark Comparison", level=1)

        paragraph = doc.add_paragraph()
        run_etf = paragraph.add_run("ETF: ")
        run_etf.bold = True
        paragraph.add_run(f"{benchmark['Name']} ({benchmark['Benchmark']})")

        benchmark_metrics = {
            "PE Ratio": benchmark.get("PE"),
            "PB Ratio": benchmark.get("PB"),
            "ROE": benchmark.get("ROE")
        }

        for label, value in benchmark_metrics.items():
            paragraph = doc.add_paragraph()
            run_label = paragraph.add_run(f"{label}: ")
            run_label.bold = True
            style = "percent" if label == "ROE" else "ratio"
            paragraph.add_run(format_number(value, style))

    # Analyst Price Targets
    if targets:
        doc.add_heading("Analyst Price Targets", level=1)
        try:
            low = float(targets.get("targetLow", 0))
            mean = float(targets.get("targetMean", 0))
            high = float(targets.get("targetHigh", 0))
            doc.add_paragraph(f"Target Range: ${low:.2f} - ${high:.2f}")
            doc.add_paragraph(f"Average Target: ${mean:.2f}")
        except:
            doc.add_paragraph("Price targets could not be processed.")

    # Valuation Label
    if valuation_label:
        doc.add_paragraph(f"Valuation: {valuation_label.capitalize()}")

    # AI Investment Summary (Structured Formatting)
    doc.add_heading("AI Investment Summary", level=1)

    if ai_summary:
        for line in clean_text(ai_summary).splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith("**") and line.endswith("**"):
                doc.add_heading(line.strip("* "), level=2)
            elif line.startswith("**") and ":**" in line:
                parts = line.strip("*").split(":**")
                heading = parts[0].strip()
                remainder = parts[1].strip() if len(parts) > 1 else ""
                paragraph = doc.add_paragraph()
                run_bold = paragraph.add_run(f"{heading}: ")
                run_bold.bold = True
                paragraph.add_run(remainder)
            else:
                doc.add_paragraph(line)
    else:
        doc.add_paragraph("Summary not available.")

    # Chart Section
    if chart_image:
        try:
            doc.add_heading("Selected Stock Price Chart", level=1)

            # Insert chart image if valid buffer
            if isinstance(chart_image, io.BytesIO):
                chart_image.seek(0)
                doc.add_picture(chart_image, width=Inches(6))

                # Label the selected time range (e.g. 6M, 1Y, etc.)
                if selected_range:
                    doc.add_paragraph(f"Time Range Selected: {selected_range}")
            else:
                doc.add_paragraph("Chart image format not recognized.")
        except Exception as e:
            doc.add_paragraph(f"Chart could not be added due to an error: {e}")

    return doc
